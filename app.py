from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file, flash
from datetime import datetime, date, time
from zoneinfo import ZoneInfo
from io import BytesIO
import json
import pandas as pd
from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, Date, Text, Float
from sqlalchemy.orm import sessionmaker, declarative_base, relationship
from docx import Document

app = Flask(__name__)
app.secret_key = "change-me-in-production"

engine = create_engine("sqlite:///school_tools.db", echo=False, future=True)
SessionLocal = sessionmaker(bind=engine, future=True)
Base = declarative_base()

TZ = ZoneInfo("Asia/Riyadh")

class Setting(Base):
    __tablename__ = "settings"
    key = Column(String, primary_key=True)
    value = Column(Text)

class Class(Base):
    __tablename__ = "classes"
    id = Column(Integer, primary_key=True)
    name = Column(String, nullable=False)
    grade = Column(String, nullable=True)
    students = relationship("Student", back_populates="class_")
    schedules = relationship("Schedule", back_populates="class_")

class Student(Base):
    __tablename__ = "students"
    id = Column(Integer, primary_key=True)
    full_name = Column(String, nullable=False, index=True)
    class_id = Column(Integer, ForeignKey("classes.id"))
    class_ = relationship("Class", back_populates="students")

class Schedule(Base):
    __tablename__ = "schedule"
    id = Column(Integer, primary_key=True)
    day_of_week = Column(Integer, nullable=False)  # 0=Sun..4=Thu
    period = Column(Integer, nullable=False)       # 1..7
    subject = Column(String, nullable=False)
    start_time = Column(String, nullable=False)    # "HH:MM"
    end_time = Column(String, nullable=False)      # "HH:MM"
    class_id = Column(Integer, ForeignKey("classes.id"))
    class_ = relationship("Class", back_populates="schedules")

class Attendance(Base):
    __tablename__ = "attendance"
    id = Column(Integer, primary_key=True)
    student_id = Column(Integer, ForeignKey("students.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    date = Column(Date, nullable=False)
    period = Column(Integer, nullable=False)
    status = Column(String, nullable=False)  # present / absent / excused

class Behavior(Base):
    __tablename__ = "behavior"
    id = Column(Integer, primary_key=True)
    student_id = Column(Integer, ForeignKey("students.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    date = Column(Date, nullable=False)
    period = Column(Integer, nullable=False)
    type = Column(String, nullable=False)  # positive / negative
    tag = Column(String, nullable=False)
    note = Column(Text, nullable=True)

class Works(Base):
    __tablename__ = "works"
    id = Column(Integer, primary_key=True)
    student_id = Column(Integer, ForeignKey("students.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    term = Column(String, default="T1")
    slots_json = Column(Text, nullable=False)  # JSON list of 12 floats/ints
    created_at = Column(Date, default=date.today)

class Homework(Base):
    __tablename__ = "homeworks"
    id = Column(Integer, primary_key=True)
    title = Column(String, nullable=False)
    max_score = Column(Float, default=100)
    assigned_date = Column(Date, default=date.today)

class HomeworkGrade(Base):
    __tablename__ = "homework_grades"
    id = Column(Integer, primary_key=True)
    homework_id = Column(Integer, ForeignKey("homeworks.id"))
    student_id = Column(Integer, ForeignKey("students.id"))
    score = Column(Float, default=0.0)

class Test(Base):
    __tablename__ = "tests"
    id = Column(Integer, primary_key=True)
    title = Column(String, nullable=False)
    max_score = Column(Float, default=100)
    test_date = Column(Date, default=date.today)

class TestGrade(Base):
    __tablename__ = "test_grades"
    id = Column(Integer, primary_key=True)
    test_id = Column(Integer, ForeignKey("tests.id"))
    student_id = Column(Integer, ForeignKey("students.id"))
    score = Column(Float, default=0.0)

Base.metadata.create_all(engine)

# Seed

def seed_defaults():
    db = SessionLocal()
    try:
        if not db.query(Setting).first():
            db.add(Setting(key="teacher_name", value="محمد أحمد الزهراني"))
            db.add(Setting(key="period_duration_minutes", value="45"))
        if not db.query(Class).first():
            c1 = Class(name="أ", grade="سادس")
            c2 = Class(name="ب", grade="سادس")
            db.add_all([c1, c2]); db.flush()
            s1 = [Student(full_name=n, class_id=c1.id) for n in ["أحمد علي", "سعد محمد", "أنس فهد"]]
            s2 = [Student(full_name=n, class_id=c2.id) for n in ["حسن عمر", "يزن خالد", "رامي سلطان"]]
            db.add_all(s1 + s2)
        if not db.query(Schedule).first():
            c1 = db.query(Class).filter_by(name="أ").first()
            base = [("08:00","08:45"),("08:50","09:35"),("09:40","10:25"),("10:40","11:25"),("11:30","12:15"),("12:20","13:05"),("13:10","13:55")]
            subs = ["علوم","رياضيات","قرآن","لغة عربية","علوم","رياضيات","تقنية"]
            for i,(t1,t2) in enumerate(base, start=1):
                db.add(Schedule(day_of_week=0, period=i, subject=subs[i-1], start_time=t1, end_time=t2, class_id=c1.id))
        db.commit()
    finally:
        db.close()

seed_defaults()

def get_setting(key, default=""):
    db = SessionLocal()
    try:
        rec = db.query(Setting).filter_by(key=key).first()
        return rec.value if rec else default
    finally:
        db.close()

def saudi_school_dow():
    d = datetime.now(TZ).weekday()  # Mon=0
    mapping = {6:0, 0:1, 1:2, 2:3, 3:4}
    return mapping.get(d, 0)

def get_todays_schedule():
    db = SessionLocal()
    try:
        dow = saudi_school_dow()
        sch = (db.query(Schedule, Class)
               .join(Class, Schedule.class_id==Class.id)
               .filter(Schedule.day_of_week==dow)
               .order_by(Schedule.period.asc())
               .all())
        res = []
        for s,c in sch:
            res.append({"id": s.id, "period": s.period, "subject": s.subject, "start_time": s.start_time, "end_time": s.end_time, "class_name": c.name, "class_id": c.id})
        return res
    finally:
        db.close()

@app.context_processor
def inject_teacher():
    return {"teacher_name": get_setting("teacher_name", "معلم العلوم")}

@app.route("/")
def index():
    schedule = get_todays_schedule()
    return render_template("index.html", schedule=schedule)

@app.route("/settings", methods=["GET","POST"]) 
def settings():
    db = SessionLocal()
    try:
        if request.method == "POST":
            tname = request.form.get("teacher_name","" ).strip()
            dur = request.form.get("period_duration_minutes","45").strip()
            for k,v in [("teacher_name", tname), ("period_duration_minutes", dur)]:
                rec = db.query(Setting).filter_by(key=k).first()
                if rec: rec.value = v
                else: db.add(Setting(key=k, value=v))
            db.commit(); flash("تم حفظ الإعدادات.", "success"); return redirect(url_for("settings"))
        return render_template("settings.html", teacher_name=get_setting("teacher_name",""), duration=get_setting("period_duration_minutes","45"))
    finally:
        db.close()

@app.route("/classes", methods=["GET"]) 
def classes():
    db = SessionLocal();
    try:
        classes = db.query(Class).all()
        return render_template("classes.html", classes=classes)
    finally:
        db.close()

@app.route("/classes/add", methods=["POST"]) 
def add_class():
    db = SessionLocal()
    try:
        name = request.form.get("name").strip(); grade = request.form.get("grade","" ).strip()
        if name:
            db.add(Class(name=name, grade=grade)); db.commit(); flash("تمت إضافة الفصل.","success")
        return redirect(url_for("classes"))
    finally:
        db.close()

@app.route("/students/<int:class_id>", methods=["GET","POST"]) 
def students(class_id):
    db = SessionLocal()
    try:
        cls = db.query(Class).get(class_id)
        if request.method == "POST":
            name = request.form.get("full_name","" ).strip()
            if name:
                db.add(Student(full_name=name, class_id=class_id)); db.commit()
        studs = db.query(Student).filter_by(class_id=class_id).order_by(Student.full_name.asc()).all()
        return render_template("students.html", cls=cls, students=studs)
    finally:
        db.close()

@app.route("/students/<int:class_id>/delete/<int:student_id>", methods=["POST"]) 
def delete_student(class_id, student_id):
    db = SessionLocal()
    try:
        st = db.query(Student).get(student_id)
        if st: db.delete(st); db.commit(); flash("تم حذف الطالب.","success")
        return redirect(url_for("students", class_id=class_id))
    finally:
        db.close()

@app.route("/students/<int:class_id>/import", methods=["POST"]) 
def import_students(class_id):
    file = request.files.get("file")
    if not file: flash("لم يتم رفع ملف.","error"); return redirect(url_for("students", class_id=class_id))
    db = SessionLocal()
    try:
        df = pd.read_excel(file, engine="openpyxl")
        if "الطالب" not in df.columns:
            flash("يجب أن يحتوي الملف على عمود باسم 'الطالب'", "error"); return redirect(url_for("students", class_id=class_id))
        count=0
        for name in df["الطالب"].dropna().astype(str).str.strip().tolist():
            if name:
                db.add(Student(full_name=name, class_id=class_id)); count+=1
        db.commit(); flash(f"تم استيراد {count} طالب.", "success")
        return redirect(url_for("students", class_id=class_id))
    finally:
        db.close()

@app.route("/api/students")
def api_students():
    class_id = request.args.get("class_id", type=int)
    q = request.args.get("q","" ).strip()
    db = SessionLocal()
    try:
        query = db.query(Student).filter(Student.class_id==class_id)
        if q and len(q) >= 3:
            query = query.filter(Student.full_name.like(f"%{q}%"))
        res = query.order_by(Student.full_name.asc()).limit(50).all()
        return jsonify([{"id":s.id, "full_name":s.full_name} for s in res])
    finally:
        db.close()

@app.route("/attendance", methods=["GET","POST"]) 
def attendance():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        selected_class_id = request.args.get("class_id", type=int) or (classes[0].id if classes else None)
        dow = saudi_school_dow()
        sch = db.query(Schedule).filter_by(day_of_week=dow, class_id=selected_class_id).order_by(Schedule.period.asc()).all()
        selected_period = request.args.get("period", type=int) or (sch[0].period if sch else 1)
        selected_date = request.args.get("date") or date.today().isoformat()
        if request.method == "POST":
            class_id = int(request.form["class_id"])
            period = int(request.form["period"])
            dt = date.fromisoformat(request.form["date"])
            for key, val in request.form.items():
                if key.startswith("status_student_"):
                    sid = int(key.split("_")[-1])
                    db.add(Attendance(student_id=sid, class_id=class_id, date=dt, period=period, status=val))
            db.commit(); flash("تم حفظ الغياب.","success")
            return redirect(url_for("attendance", class_id=class_id, period=period, date=dt.isoformat()))
        students = db.query(Student).filter_by(class_id=selected_class_id).order_by(Student.full_name.asc()).all()
        return render_template("attendance.html", classes=classes, students=students, schedule=sch,
                               selected_class_id=selected_class_id, selected_period=selected_period, selected_date=selected_date)
    finally:
        db.close()

@app.route("/behavior", methods=["GET","POST"]) 
def behavior():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        selected_class_id = request.args.get("class_id", type=int) or (classes[0].id if classes else None)
        dow = saudi_school_dow()
        sch = db.query(Schedule).filter_by(day_of_week=dow, class_id=selected_class_id).order_by(Schedule.period.asc()).all()
        selected_period = request.args.get("period", type=int) or (sch[0].period if sch else 1)
        selected_date = request.args.get("date") or date.today().isoformat()
        if request.method == "POST":
            class_id = int(request.form["class_id"])
            period = int(request.form["period"])
            dt = date.fromisoformat(request.form["date"])
            student_id = int(request.form["student_id"])
            btype = request.form["type"]
            tag = request.form["tag"].strip()
            note = request.form.get("note","" ).strip() or None
            db.add(Behavior(student_id=student_id, class_id=class_id, date=dt, period=period, type=btype, tag=tag, note=note))
            db.commit(); flash("تم حفظ السلوك.","success")
            return redirect(url_for("behavior", class_id=class_id, period=period, date=dt.isoformat()))
        return render_template("behavior.html", classes=classes, schedule=sch,
                               selected_class_id=selected_class_id, selected_period=selected_period, selected_date=selected_date)
    finally:
        db.close()

@app.route("/works", methods=["GET","POST"]) 
def works():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        selected_class_id = request.args.get("class_id", type=int) or (classes[0].id if classes else None)
        term = request.args.get("term","T1")
        students = db.query(Student).filter_by(class_id=selected_class_id).order_by(Student.full_name.asc()).all()
        if request.method == "POST":
            term = request.form.get("term","T1")
            for s in students:
                slots = []
                for i in range(1,13):
                    v = request.form.get(f"slot_{s.id}_{i}", "").strip()
                    slots.append(float(v) if v!='' else 0.0)
                db.add(Works(student_id=s.id, class_id=selected_class_id, term=term, slots_json=json.dumps(slots)))
            db.commit(); flash("تم حفظ الأعمال الأدائية.","success")
            return redirect(url_for("works", class_id=selected_class_id, term=term))
        latest = {}
        for s in students:
            rec = db.query(Works).filter_by(student_id=s.id, class_id=selected_class_id, term=term).order_by(Works.id.desc()).first()
            latest[s.id] = json.loads(rec.slots_json) if rec else [0]*12
        return render_template("works.html", classes=classes, students=students, selected_class_id=selected_class_id, term=term, latest=latest)
    finally:
        db.close()

@app.route("/homeworks", methods=["GET"]) 
def homeworks():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        selected_class_id = request.args.get("class_id", type=int) or (classes[0].id if classes else None)
        students = db.query(Student).filter_by(class_id=selected_class_id).order_by(Student.full_name.asc()).all()
        homeworks = db.query(Homework).order_by(Homework.assigned_date.asc()).all()
        # build grades map
        gmap = {}
        hids = [h.id for h in homeworks]
        if hids:
            rows = db.query(HomeworkGrade).filter(HomeworkGrade.student_id.in_([s.id for s in students]), HomeworkGrade.homework_id.in_(hids)).all()
            for r in rows:
                gmap[f"{r.student_id}_{r.homework_id}"] = r.score
        return render_template("homeworks.html", classes=classes, selected_class_id=selected_class_id, students=students, homeworks=homeworks, grades=gmap, today=date.today().isoformat())
    finally:
        db.close()

@app.route("/homeworks/add", methods=["POST"]) 
def add_homework():
    db = SessionLocal()
    try:
        title = request.form.get("title").strip()
        max_score = float(request.form.get("max_score", 100))
        assigned_date = date.fromisoformat(request.form.get("assigned_date"))
        if title:
            db.add(Homework(title=title, max_score=max_score, assigned_date=assigned_date)); db.commit(); flash("تمت إضافة الواجب.","success")
        return redirect(url_for("homeworks", class_id=int(request.form.get("class_id"))))
    finally:
        db.close()

@app.route("/homeworks/save", methods=["POST"]) 
def save_homework_grades():
    db = SessionLocal()
    try:
        class_id = int(request.form.get("class_id"))
        for k,v in request.form.items():
            if k.startswith("grade_") and v.strip()!='':
                _, sid, hid = k.split("_")
                sid, hid = int(sid), int(hid)
                score = float(v)
                # upsert
                rec = db.query(HomeworkGrade).filter_by(student_id=sid, homework_id=hid).first()
                if rec: rec.score = score
                else: db.add(HomeworkGrade(student_id=sid, homework_id=hid, score=score))
        db.commit(); flash("تم حفظ درجات الواجبات.","success")
        return redirect(url_for("homeworks", class_id=class_id))
    finally:
        db.close()

@app.route("/homeworks/import/<int:class_id>", methods=["POST"]) 
def import_homework_excel(class_id):
    file = request.files.get("file")
    if not file:
        flash("لم يتم رفع ملف.","error"); return redirect(url_for("homeworks", class_id=class_id))
    db = SessionLocal()
    try:
        df = pd.read_excel(file, sheet_name="homeworks", engine="openpyxl")
        if "الطالب" not in df.columns:
            flash("ورقة homeworks يجب أن تحتوي عمود 'الطالب'", "error"); return redirect(url_for("homeworks", class_id=class_id))
        # Ensure homeworks exist by title
        titles = [c for c in df.columns if c != "الطالب"]
        existing = {h.title: h for h in db.query(Homework).filter(Homework.title.in_(titles)).all()}
        for t in titles:
            if t not in existing:
                h = Homework(title=t, max_score=100, assigned_date=date.today())
                db.add(h); db.flush(); existing[t]=h
        # Map students in this class
        studs = {s.full_name: s.id for s in db.query(Student).filter_by(class_id=class_id).all()}
        n=0
        for _,row in df.iterrows():
            name = str(row["الطالب"]).strip()
            sid = studs.get(name)
            if not sid: continue
            for t in titles:
                val = row.get(t)
                if pd.isna(val):
                    continue
                hid = existing[t].id
                rec = db.query(HomeworkGrade).filter_by(student_id=sid, homework_id=hid).first()
                if rec: rec.score = float(val)
                else: db.add(HomeworkGrade(student_id=sid, homework_id=hid, score=float(val)))
                n+=1
        db.commit(); flash(f"تم استيراد/تحديث {n} درجة واجب.","success")
        return redirect(url_for("homeworks", class_id=class_id))
    finally:
        db.close()

@app.route("/tests", methods=["GET"]) 
def tests():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        selected_class_id = request.args.get("class_id", type=int) or (classes[0].id if classes else None)
        students = db.query(Student).filter_by(class_id=selected_class_id).order_by(Student.full_name.asc()).all()
        tests = db.query(Test).order_by(Test.test_date.asc()).all()
        gmap = {}
        tids = [t.id for t in tests]
        if tids:
            rows = db.query(TestGrade).filter(TestGrade.student_id.in_([s.id for s in students]), TestGrade.test_id.in_(tids)).all()
            for r in rows:
                gmap[f"{r.student_id}_{r.test_id}"] = r.score
        return render_template("tests.html", classes=classes, selected_class_id=selected_class_id, students=students, tests=tests, grades=gmap, today=date.today().isoformat())
    finally:
        db.close()

@app.route("/tests/add", methods=["POST"]) 
def add_test():
    db = SessionLocal()
    try:
        title = request.form.get("title").strip()
        max_score = float(request.form.get("max_score", 100))
        test_date = date.fromisoformat(request.form.get("test_date"))
        if title:
            db.add(Test(title=title, max_score=max_score, test_date=test_date)); db.commit(); flash("تمت إضافة الاختبار.","success")
        return redirect(url_for("tests", class_id=int(request.form.get("class_id"))))
    finally:
        db.close()

@app.route("/tests/save", methods=["POST"]) 
def save_test_grades():
    db = SessionLocal()
    try:
        class_id = int(request.form.get("class_id"))
        for k,v in request.form.items():
            if k.startswith("grade_") and v.strip()!='':
                _, sid, tid = k.split("_")
                sid, tid = int(sid), int(tid)
                score = float(v)
                rec = db.query(TestGrade).filter_by(student_id=sid, test_id=tid).first()
                if rec: rec.score = score
                else: db.add(TestGrade(student_id=sid, test_id=tid, score=score))
        db.commit(); flash("تم حفظ درجات الاختبارات.","success")
        return redirect(url_for("tests", class_id=class_id))
    finally:
        db.close()

@app.route("/tests/import/<int:class_id>", methods=["POST"]) 
def import_test_excel(class_id):
    file = request.files.get("file")
    if not file:
        flash("لم يتم رفع ملف.","error"); return redirect(url_for("tests", class_id=class_id))
    db = SessionLocal()
    try:
        df = pd.read_excel(file, sheet_name="tests", engine="openpyxl")
        if "الطالب" not in df.columns:
            flash("ورقة tests يجب أن تحتوي عمود 'الطالب'", "error"); return redirect(url_for("tests", class_id=class_id))
        titles = [c for c in df.columns if c != "الطالب"]
        existing = {t.title: t for t in db.query(Test).filter(Test.title.in_(titles)).all()}
        for t in titles:
            if t not in existing:
                tt = Test(title=t, max_score=100, test_date=date.today())
                db.add(tt); db.flush(); existing[t]=tt
        studs = {s.full_name: s.id for s in db.query(Student).filter_by(class_id=class_id).all()}
        n=0
        for _,row in df.iterrows():
            name = str(row["الطالب"]).strip()
            sid = studs.get(name)
            if not sid: continue
            for t in titles:
                val = row.get(t)
                if pd.isna(val): continue
                tid = existing[t].id
                rec = db.query(TestGrade).filter_by(student_id=sid, test_id=tid).first()
                if rec: rec.score = float(val)
                else: db.add(TestGrade(student_id=sid, test_id=tid, score=float(val)))
                n+=1
        db.commit(); flash(f"تم استيراد/تحديث {n} درجة اختبار.","success")
        return redirect(url_for("tests", class_id=class_id))
    finally:
        db.close()

@app.route("/report/student/<int:student_id>") 
def report_student(student_id):
    db = SessionLocal()
    try:
        s = db.query(Student).get(student_id)
        att = db.query(Attendance).filter(Attendance.student_id==student_id).all()
        absences = [a for a in att if a.status=="absent"]
        beh = db.query(Behavior).filter(Behavior.student_id==student_id).all()
        pos = [b for b in beh if b.type=="positive"]
        neg = [b for b in beh if b.type=="negative"]
        w = db.query(Works).filter(Works.student_id==student_id).order_by(Works.id.desc()).first()
        slots = json.loads(w.slots_json) if w else [0]*12
        works_count = sum(1 for x in slots if x>0)
        works_avg = round(sum(slots)/12.0, 2) if slots else 0.0
        hgs = db.query(HomeworkGrade).filter(HomeworkGrade.student_id==student_id).all()
        hw_count = len(hgs); hw_avg = round(sum([g.score for g in hgs])/hw_count, 2) if hw_count>0 else 0.0
        tgs = db.query(TestGrade).filter(TestGrade.student_id==student_id).all()
        test_count = len(tgs); test_avg = round(sum([g.score for g in tgs])/test_count, 2) if test_count>0 else 0.0
        teacher_name = get_setting("teacher_name","معلم العلوم")
        return render_template("report_student.html", s=s, absences=absences, pos=pos, neg=neg, slots=slots, works_count=works_count, works_avg=works_avg, hw_count=hw_count, hw_avg=hw_avg, test_count=test_count, test_avg=test_avg, teacher_name=teacher_name, today=date.today())
    finally:
        db.close()

@app.route("/report/class/<int:class_id>") 
def report_class(class_id):
    db = SessionLocal()
    try:
        cls = db.query(Class).get(class_id)
        students = db.query(Student).filter_by(class_id=class_id).order_by(Student.full_name.asc()).all()
        rows = []
        total_absences = 0
        abs_count_map = {}
        for s in students:
            att = db.query(Attendance).filter(Attendance.student_id==s.id).all()
            absences = len([a for a in att if a.status=="absent"])
            total_absences += absences
            abs_count_map[s.full_name] = absences
            beh = db.query(Behavior).filter(Behavior.student_id==s.id).all()
            pos = len([b for b in beh if b.type=="positive"])
            neg = len([b for b in beh if b.type=="negative"])
            w = db.query(Works).filter(Works.student_id==s.id).order_by(Works.id.desc()).first()
            slots = json.loads(w.slots_json) if w else [0]*12
            works_count = sum(1 for x in slots if x>0)
            works_avg = round(sum(slots)/12.0, 2) if slots else 0.0
            hgs = db.query(HomeworkGrade).filter(HomeworkGrade.student_id==s.id).all()
            hw_avg = round(sum([g.score for g in hgs])/len(hgs), 2) if hgs else 0.0
            tgs = db.query(TestGrade).filter(TestGrade.student_id==s.id).all()
            test_avg = round(sum([g.score for g in tgs])/len(tgs), 2) if tgs else 0.0
            rows.append({"name": s.full_name, "absences": absences, "pos": pos, "neg": neg, "works_count": works_count, "works_avg": works_avg, "hw_avg": hw_avg, "test_avg": test_avg})
        top_absent_name = max(abs_count_map, key=abs_count_map.get) if abs_count_map else "—"
        top_absent_count = abs_count_map.get(top_absent_name, 0)
        return render_template("report_class.html", cls=cls, students=students, rows=rows, total_absences=total_absences, top_absent_name=top_absent_name, top_absent_count=top_absent_count)
    finally:
        db.close()

@app.route("/reports") 
def reports_all():
    db = SessionLocal();
    try:
        classes = db.query(Class).order_by(Class.name.asc()).all()
        return render_template("report_all.html", classes=classes)
    finally:
        db.close()

@app.route("/export/excel/class/<int:class_id>") 
def export_excel_class(class_id):
    db = SessionLocal()
    try:
        cls = db.query(Class).get(class_id)
        students = db.query(Student).filter_by(class_id=class_id).order_by(Student.full_name.asc()).all()
        att = db.query(Attendance).filter(Attendance.class_id==class_id).all()
        st_map = {st.id: st.full_name for st in students}
        att_rows = [{"الطالب": st_map.get(a.student_id, "—"), "التاريخ": a.date.isoformat(), "الحصة": a.period, "الحالة": "حاضر" if a.status=="present" else ("غائب" if a.status=="absent" else "مُعذّر")} for a in att]
        beh = db.query(Behavior).filter(Behavior.class_id==class_id).all()
        beh_rows = [{"الطالب": st_map.get(b.student_id, "—"), "التاريخ": b.date.isoformat(), "الحصة": b.period, "النوع": "إيجابي" if b.type=="positive" else "سلبي", "السلوك": b.tag, "ملاحظة": b.note or ""} for b in beh]
        works_latest = []
        for st in students:
            w = db.query(Works).filter_by(student_id=st.id, class_id=class_id).order_by(Works.id.desc()).first()
            row = {"الطالب": st.full_name}
            if w:
                slots = json.loads(w.slots_json)
                for i,val in enumerate(slots, start=1): row[f"عمل {i}"] = val
                row["عدد المسلّم"] = sum(1 for x in slots if x>0)
                row["متوسط"] = round(sum(slots)/12.0, 2)
            else:
                for i in range(1,13): row[f"عمل {i}"] = 0
                row["عدد المسلّم"] = 0; row["متوسط"] = 0
            works_latest.append(row)
        bio = BytesIO()
        with pd.ExcelWriter(bio, engine="openpyxl") as writer:
            pd.DataFrame([{"اسم الفصل": cls.name, "الصف": cls.grade}]).to_excel(writer, sheet_name="الفصل", index=False)
            pd.DataFrame(att_rows).to_excel(writer, sheet_name="الغياب", index=False)
            pd.DataFrame(beh_rows).to_excel(writer, sheet_name="السلوك", index=False)
            pd.DataFrame(works_latest).to_excel(writer, sheet_name="الأعمال الأدائية", index=False)
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=f"class_{cls.name}_export.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    finally:
        db.close()

@app.route("/export/word/student/<int:student_id>") 
def export_word_student(student_id):
    db = SessionLocal()
    try:
        s = db.query(Student).get(student_id)
        doc = Document()
        teacher_name = get_setting("teacher_name","معلم العلوم")
        doc.add_heading(f"تقرير الطالب: {s.full_name}", 0)
        doc.add_paragraph(f"الفصل: {s.class_.name} — الصف: {s.class_.grade}")
        att = db.query(Attendance).filter(Attendance.student_id==student_id).all()
        absences = [a for a in att if a.status=="absent"]
        doc.add_heading("الغياب", level=1)
        doc.add_paragraph(f"عدد أيام الغياب: {len(absences)}")
        for a in absences[:100]:
            doc.add_paragraph(f"- {a.date.isoformat()} — حصة {a.period}")
        beh = db.query(Behavior).filter(Behavior.student_id==student_id).all()
        pos = [b for b in beh if b.type=="positive"]
        neg = [b for b in beh if b.type=="negative"]
        doc.add_heading("السلوك الإيجابي", level=1)
        for b in pos[:100]: doc.add_paragraph(f"- {b.date.isoformat()} — {b.tag} (حصة {b.period}) {('— ' + b.note) if b.note else ''}")
        doc.add_heading("السلوك السلبي", level=1)
        for b in neg[:100]: doc.add_paragraph(f"- {b.date.isoformat()} — {b.tag} (حصة {b.period}) {('— ' + b.note) if b.note else ''}")
        w = db.query(Works).filter(Works.student_id==student_id).order_by(Works.id.desc()).first()
        doc.add_heading("الأعمال الأدائية", level=1)
        if w:
            slots = json.loads(w.slots_json)
            doc.add_paragraph(f"عدد الأعمال المسلّمة: {sum(1 for x in slots if x>0)} — المتوسط: {round(sum(slots)/12.0,2)}")
        else:
            doc.add_paragraph("لا توجد بيانات أعمال أدائية.")
        hgs = db.query(HomeworkGrade).filter(HomeworkGrade.student_id==student_id).all()
        tgs = db.query(TestGrade).filter(TestGrade.student_id==student_id).all()
        doc.add_heading("الواجبات", level=1)
        if hgs:
            avg = round(sum([g.score for g in hgs])/len(hgs),2)
            doc.add_paragraph(f"عدد الواجبات: {len(hgs)} — المتوسط: {avg}")
        else:
            doc.add_paragraph("لا توجد بيانات واجبات.")
        doc.add_heading("الاختبارات", level=1)
        if tgs:
            avg = round(sum([g.score for g in tgs])/len(tgs),2)
            doc.add_paragraph(f"عدد الاختبارات: {len(tgs)} — المتوسط: {avg}")
        else:
            doc.add_paragraph("لا توجد بيانات اختبارات.")
        doc.add_paragraph("\n"); doc.add_paragraph(f"معلم العلوم: {teacher_name}")
        bio = BytesIO(); doc.save(bio); bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=f"report_{s.full_name}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    finally:
        db.close()

@app.route("/schedule", methods=["GET"]) 
def schedule_page():
    db = SessionLocal()
    try:
        classes = db.query(Class).all()
        sch = db.query(Schedule).order_by(Schedule.day_of_week.asc(), Schedule.period.asc()).all()
        return render_template("schedule.html", classes=classes, schedule=sch)
    finally:
        db.close()

@app.route("/schedule/add", methods=["POST"]) 
def add_schedule():
    db = SessionLocal()
    try:
        day_of_week = int(request.form.get("day_of_week"))
        class_id = int(request.form.get("class_id"))
        period = int(request.form.get("period"))
        subject = request.form.get("subject").strip()
        start_time = request.form.get("start_time").strip()
        end_time = request.form.get("end_time").strip()
        db.add(Schedule(day_of_week=day_of_week, class_id=class_id, period=period, subject=subject, start_time=start_time, end_time=end_time))
        db.commit(); flash("تمت إضافة سطر إلى الجدول.","success")
        return redirect(url_for("schedule_page"))
    finally:
        db.close()

@app.route("/schedule/delete/<int:sid>", methods=["POST"]) 
def delete_schedule(sid):
    db = SessionLocal()
    try:
        rec = db.query(Schedule).get(sid)
        if rec: db.delete(rec); db.commit(); flash("تم الحذف.","success")
        return redirect(url_for("schedule_page"))
    finally:
        db.close()

@app.route("/api/today") 
def api_today():
    return jsonify(get_todays_schedule())

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080, debug=True)
